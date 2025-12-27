import re
import os
import json
import logging
import spacy
import cv2
import pytesseract
from pytesseract import Output
import fitz  # PyMuPDF
import docx
from docx.shared import RGBColor
from PIL import Image
import requests
import google.generativeai as genai
from steganography.steganography import Steganography

pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# --- Lazy Loading for AI Models ---
transformers = None
torch = None
local_ner_pipeline = None

# --- INITIALIZATION ---
logger = logging.getLogger(__name__)
nlp_spacy = spacy.load("en_core_web_sm")
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
GEMINI_MODEL_NAME = "gemini-2.5-flash"
REDACTION_TEXT = "[REDACTED]"

# --- EXPANDED PII PATTERNS ---
PII_PATTERNS = {
    'EMAIL': re.compile(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'),
    'PHONE': re.compile(r'(\+91[\-\s]?)?[0]?(91)?[6-9]\d{9}'),
    'CREDIT_CARD': re.compile(r'\b(?:\d[ -]*?){13,16}\b'),
    'AADHAAR': re.compile(r'\b\d{4}[ -]?\d{4}[ -]?\d{4}\b'),
    'PAN': re.compile(r'[A-Z]{5}[0-9]{4}[A-Z]{1}'),
    'UPI_ID': re.compile(r'[a-zA-Z0-9.\-_]{2,256}@[a-zA-Z]{2,64}'),
    'VOTER_ID': re.compile(r'[A-Z]{3}[0-9]{7}'),
    'DRIVING_LICENCE': re.compile(r'[A-Z]{2}[0-9]{2}\s?[0-9]{4}\s?[0-9]{7}'), # Common DL format
    'GSTIN': re.compile(r'[0-9]{2}[A-Z]{5}[0-9]{4}[A-Z]{1}[1-9A-Z]{1}Z[0-9A-Z]{1}')
}

# --- LOCAL AI MODEL FUNCTION ---
def initialize_local_ner():
    """Initializes the local Hugging Face NER model."""
    global transformers, torch, local_ner_pipeline
    if local_ner_pipeline is None:
        try:
            transformers = __import__('transformers')
            torch = __import__('torch')
            # Switched to a smaller, faster "distilled" model for better CPU performance.
            model_name = "distilbert-base-cased-distilled-squad"
            logger.info(f"Initializing local NER model: {model_name}. This may take a moment...")
            
            local_ner_pipeline = transformers.pipeline(
                "ner", 
                model=model_name, 
                tokenizer=model_name, 
                aggregation_strategy="simple"
            )
            logger.info("Local NER model initialized successfully.")
        except ImportError:
            logger.error("Hugging Face 'transformers' or 'torch' library not found. Please install them.")
            raise
        except Exception as e:
            logger.error(f"Failed to load local NER model: {e}")
            raise

def detect_pii_with_local_ner(text):
    if local_ner_pipeline is None: initialize_local_ner()
    if not text or not text.strip() or not local_ner_pipeline: return []
    try:
        ner_results = local_ner_pipeline(text)
        pii_entities = set()
        for entity in ner_results:
            if entity['entity_group'].upper() in ['PER', 'LOC', 'ORG', 'MISC', 'ANSWER']:
                pii_entities.add(entity['word'])
        logger.info(f"Local model detected {len(pii_entities)} PII elements.")
        return list(pii_entities)
    except Exception as e:
        logger.error(f"Error during local NER processing: {e}")
        return []

def detect_pii_with_gemini(text, api_key):
    if not text or not text.strip(): return []
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(GEMINI_MODEL_NAME)
        prompt = f"""Analyze the following text to identify PII (Personally Identifiable Information). 
        Return ONLY a valid JSON array of strings containing the exact PII text found. 
        Do not include any other text, markdown formatting, or code blocks.
        Example: ["John Doe", "john.doe@email.com", "123-456-7890"]

        Text to analyze:
        {text}
        """
        response = model.generate_content(prompt)
        content = response.text
        # Clean up potential markdown formatting if the model disregards instructions
        content = content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        
        json_str = content.strip()
        return [str(item) for item in json.loads(json_str)]
    except Exception as e:
        logger.error(f"Gemini API call failed: {e}")
        # Fallback to empty list so we don't crash the whole process, but log error
        return []



def find_all_pii(text, options):
    logger.info(f"Processing with options: {options}")
    all_pii_text = set()
    if options.get('ai_model_choice') == 'local':
        all_pii_text.update(detect_pii_with_local_ner(text))
    elif options.get('gemini_api_key'):
        all_pii_text.update(detect_pii_with_gemini(text, options['gemini_api_key']))

    for pii_type in options.get('pii_types', []):
        if pii_type in PII_PATTERNS:
            for match in PII_PATTERNS[pii_type].finditer(text):
                all_pii_text.add(match.group(0))
    if 'NAME' in options.get('pii_types', []):
        doc = nlp_spacy(text)
        for ent in doc.ents:
            if ent.label_ == "PERSON":
                all_pii_text.add(ent.text)
    return sorted(list(all_pii_text), key=len, reverse=True)

def apply_pdf_watermark(page):
    r = page.rect
    text = "REDACTED"
    # Calculate font size to span a good portion of the page diagonally
    font_size = min(r.width, r.height) / 6
    
    # Calculate center position
    x = r.width / 2
    y = r.height / 2
    
    # Insert text with rotation. 
    # Note: insert_text expects the point to be the bottom-left of the text usually, 
    # but when rotating around a point, it's often the pivot.
    # PyMuPDF's insert_text supports 'rotate' in newer versions. 
    # If not, we use a different approach. But assuming standard PyMuPDF:
    try:
        # Use simple int for rotate if 0, 90, 180, 270. For 45, use morph.
        # Create a matrix for 45 degree rotation
        mat = fitz.Matrix(-45)
        # Point to rotate around (center of page)
        p = fitz.Point(x, y)
        
        # To center the text effectively while rotating:
        # We want the center of the text to be at 'p'.
        # insert_text places the start of the baseline at the point.
        # We need to shift the point "backwards" by half the text width and "down" by half height (conceptually)
        # before rotation. But since we rotate around 'p', it's easier to just offset the
        # insertion point relative to p.
        
        # Estimate text width (PyMuPDF's get_text_length is accurate for standard fonts)
        text_len = fitz.get_text_length(text, fontname="hebo", fontsize=font_size)
        
        # Shift start point: left by half width, down by half font size (approx vertical center)
        # Note: In PDF coords, y increases downwards.
        # We define a new point relative to P (center) that represents the bottom-left of the text
        # if it were unrotated, such that its center aligns with P.
        start_x = x - (text_len / 2)
        start_y = y + (font_size / 3) # Approximate vertical centering for cap height
        
        p_start = fitz.Point(start_x, start_y)
        
        # When using morph with a pivot (like via fitz.Matrix.rotate or manually),
        # we usually just apply the matrix to the point or let the library handle it.
        # Here we use morph=(pivot, matrix). The text is drawn at p_start, then the separate parts
        # are transformed.
        # Actually, simpler approach with morph=(px, py, matrix) or similar might be tricky.
        # Let's trust standard insert_text with morph=(p_start, matrix) effectively rotates 
        # the text AS IF it was at p_start, around p_start? No, usually morph rotates around (0,0) unless specified.
        # PyMuPDF docs say: morph=(fixed_point, matrix). "The text is rotated ... around fixed_point".
        # So if we put text at p_start, and rotate around p_start, it just spins in place (off center).
        # We want to rotate around the PAGE CENTER.
        
        morph = (p, mat) # Rotate around page center
        
        # Insert text at the calculated start point
        page.insert_text(p_start, text, 
                         fontsize=font_size, 
                         fontname="hebo", 
                         color=(0.2, 0.2, 0.2), 
                         fill_opacity=0.6, 
                         morph=morph)
    except Exception as e:
        logger.error(f"Failed to apply watermark via insert_text: {e}")
        pass

def redact_image(file_path, options):
    output_path = file_path.replace(os.path.splitext(file_path)[1], "_redacted.png")
    img_cv = cv2.cvtColor(cv2.imread(file_path), cv2.COLOR_BGR2RGB)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_RGB2GRAY)
    faces = face_cascade.detectMultiScale(gray, 1.1, 4)
    for (x, y, w, h) in faces:
        cv2.rectangle(img_cv, (x, y), (x + w, y + h), (0, 0, 0), -1)
    ocr_data = pytesseract.image_to_data(img_cv, output_type=Output.DICT)
    full_text = " ".join(ocr_data['text'])
    pii_to_redact = find_all_pii(full_text, options)
    for pii in pii_to_redact:
        for i, text in enumerate(ocr_data['text']):
            if pii in text:
                (x, y, w, h) = (ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i])
                cv2.rectangle(img_cv, (x, y), (x + w, y + h), (0, 0, 0), -1)
    
    if options.get('apply_watermark'):
        logger.info("Applying watermark to image...")
        h, w, _ = img_cv.shape
        text = "REDACTED"
        font_scale = min(w, h) / 300
        thickness = 5
        font = cv2.FONT_HERSHEY_SIMPLEX
        
        # Get text size
        (text_w, text_h), baseline = cv2.getTextSize(text, font, font_scale, thickness)
        
        # Calculate center coordinates
        text_x = int((w - text_w) / 2)
        text_y = int((h + text_h) / 2)
        
        cv2.putText(img_cv, text, (text_x, text_y), font, font_scale, (255, 0, 0), thickness, cv2.LINE_AA)
        
    final_img_pil = Image.fromarray(img_cv)
    if options.get('covert_redaction'):
        # For covert redaction on images, we hide PII in metadata/steganography but still produce an image.
        # If user ASKED for watermark, we should retain the watermarked image as the carrier? 
        # Usually steganography takes the clean image or redacted image.
        # Here we re-read the file to do OCR? No, let's use the image we just prepared.
        # Wait, the original code did `Steganography.encode(file_path, output_path, secret_message)`.
        # `file_path` is the ORIGINAL file.
        # If we overwrite `output_path` with steganography result, we lose the visual redactions/watermark we just did in `img_cv`.
        # This logic is a bit flawed for combining visual + covert.
        # Assuming covert redaction suppresses visual redaction (white/black bars), but the user just asked for watermark + covert.
        # If covert is meant to be HIDDEN, usually you don't watermark it.
        # But if we must support both:
        # 1. Take result with visual redactions (if any) and watermark.
        # 2. Encode secret message into THAT.
        
        # Current logic:
        # ocr_data_covert = pytesseract.image_to_data(file_path...
        # ...
        # Steganography.encode(file_path, output_path, secret_message)
        # This IGNORES `final_img_pil` (the one with black boxes and watermark).
        
        # Fix: Save `final_img_pil` to `output_path` first, then use THAT as input for steganography.
        final_img_pil.save(output_path, format='PNG')
        
        ocr_data_covert = pytesseract.image_to_data(final_img_pil, output_type=Output.DICT) # Use modified image for consistency? Or original?
        # Typically you safeguard the PII from the original.
        full_text_covert = " ".join(ocr_data['text']) # Reuse original OCR data for detecting what to hide?
        pii_to_hide = find_all_pii(full_text_covert, options)
        secret_message = "Redacted PII: " + ", ".join(pii_to_hide)
        
        # Encode into the already modified image
        Steganography.encode(output_path, output_path, secret_message)
        logger.info(f"Covertly hid {len(pii_to_hide)} PII elements in the image.")
        
        # Reload for metadata wiping if needed
        final_img_pil = Image.open(output_path)
    if options.get('wipe_metadata'):
        data = list(final_img_pil.getdata())
        image_without_metadata = Image.new(final_img_pil.mode, final_img_pil.size)
        image_without_metadata.putdata(data)
        image_without_metadata.save(output_path, format='PNG')
    else:
        final_img_pil.save(output_path, format='PNG')
    return output_path

def redact_pdf(file_path, options):
    doc = fitz.open(file_path)
    if options.get('wipe_metadata'): doc.set_metadata({})
    for page in doc:
        text = page.get_text("text")
        pii_to_redact = find_all_pii(text, options)
        for pii in pii_to_redact:
            areas = page.search_for(pii)
            for inst in areas:
                if options.get('covert_redaction'):
                    page.draw_rect(inst, color=(1, 1, 1), fill=(1, 1, 1), overlay=True)
                else:
                    page.add_redact_annot(inst, fill=(0, 0, 0))
        if not options.get('covert_redaction'): page.apply_redactions()
        if options.get('apply_watermark'):
            apply_pdf_watermark(page)
    output_path = file_path.replace(".pdf", "_redacted.pdf")
    doc.save(output_path, garbage=4, deflate=True, clean=True)
    doc.close()
    return output_path

def redact_docx(file_path, options):
    doc = docx.Document(file_path)
    full_text = "\n".join([p.text for p in doc.paragraphs])
    pii_to_redact = find_all_pii(full_text, options)
    for pii in pii_to_redact:
        for para in doc.paragraphs:
            if pii in para.text:
                if options.get('covert_redaction'):
                    para.text = para.text.replace(pii, ' ' * len(pii))
                else:
                    para.text = para.text.replace(pii, REDACTION_TEXT)
    output_path = file_path.replace(".docx", "_redacted.docx")
    if options.get('wipe_metadata'):
        doc.core_properties.author = None
        doc.core_properties.last_modified_by = None
    doc.save(output_path)
    return output_path

def process_file(file_path, options):
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()
    if extension in ['.png', '.jpg', '.jpeg']:
        return redact_image(file_path, options)
    elif extension == '.pdf':
        return redact_pdf(file_path, options)
    elif extension == '.docx':
        return redact_docx(file_path, options)
    else:
        raise ValueError(f"Unsupported file type: {extension}")

